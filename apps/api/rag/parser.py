from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    metadata: dict[str, str]


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return ParsedDocument(content.decode("utf-8"), {"source": filename})
    if suffix == ".pdf":
        from pypdf import PdfReader

        pages = [page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages]
        return ParsedDocument("\n\n".join(pages), {"source": filename})
    if suffix == ".docx":
        from docx import Document

        paragraphs = [paragraph.text for paragraph in Document(BytesIO(content)).paragraphs]
        return ParsedDocument("\n".join(paragraphs), {"source": filename})
    raise ValueError(f"Unsupported document type: {suffix or 'unknown'}")

